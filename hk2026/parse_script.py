#!/usr/bin/env python3
"""
parse_script.py
香港 日本産食肉加工品 商談会/交流会 (2026-08-12)
MC構成台本 (.docx) → data/hk2026.json 変換スクリプト

原稿フォーマット:
  - 1テーブル / 2カラム（時間 | 司会者進行）
  - 【…】だけの行 = セクション見出し（結合セル）
  - 台本セル内は「日：」「中：」で言語を切り替え
  - 企業紹介セクションのみ prefix 無しで日本語行/中国語行が交互

使用方法: python3 parse_script.py
"""
import docx
import json
import re
from pathlib import Path

BASE_DIR   = Path(__file__).parent.parent
DOCX_PATH  = BASE_DIR / "台本" / "2026香港商談会交流会_MC構成台本.docx"
OUT_PATH    = Path(__file__).parent / "data" / "hk2026.json"
OUT_JS_PATH = Path(__file__).parent / "data" / "hk2026.js"

KANA = re.compile(r"[぀-ゟ゠-ヺー]")

# セクション見出しの中国語訳（原稿本文の語彙に合わせて作成）
SECTION_ZH = {
    "スタート3分前のアナウンス": "開始前3分鐘廣播",
    "オープニング":             "開場致辭",
    "主催者挨拶":               "主辦方致辭",
    "企業様の紹介":             "參展企業介紹",
    "試食・商談・ネットワーキング": "試食・商談・交流",
    "終了のご挨拶":             "閉幕致辭",
}

EVENT = {
    "id":       "hk2026",
    "title":    "日本産食肉加工品 商談会 / 交流会",
    "titleZh":  "日本加工肉產品商談交流會",
    "subtitle": "HONG KONG 2026",
    "date":     "2026年8月12日（水）",
    "dateZh":   "2026年8月12日（三）",
    "timezone": "Asia/Hong_Kong",
    "tzLabel":  "HKT",
    "venue":    "Regal Kowloon Hotel 3階（71 Mody Road, Tsim Sha Tsui East, Kowloon）",
    "venueZh":  "富豪九龍酒店 3樓（尖沙咀東部麼地道71號）",
    "host":     "日本ハム・ソーセージ工業協同組合",
    "operator": "ブラボーワークス株式会社",
    "mc1": {"name": "保呂田", "fullName": "保呂田（Bravoworks）", "lang": "日本語",
            "nameZh": "保呂田", "langZh": "日語"},
    "mc2": {"name": "Cindy",  "fullName": "Cindy（Umai Communications）", "lang": "中文",
            "nameZh": "Cindy", "langZh": "中文"},
    # 台本上の実際の進行時刻（ヘッダー表記の 19:00-21:00 とは別。README参照）
    "startTime": "18:42",
    "endTime":   "20:45",
    "footnote":  "※ 時間はおおよその目安です。当日の進行状況に応じて調整してください。",
}


def is_ja(text):
    """かな（ひらがな・カタカナ）を含めば日本語行とみなす"""
    return bool(KANA.search(text))


def cell_lines(cell):
    return [p.text.strip() for p in cell.paragraphs if p.text.strip()]


def parse_roster(lines):
    """企業紹介セクション: 番号付きの日本語行/中国語行のペアを抽出"""
    intro_ja, intro_zh, roster = None, None, []
    pending = {}

    for line in lines:
        m = re.match(r"^(\d+)[.．、]\s*(.+)$", line)
        if m:
            no, body = int(m.group(1)), m.group(2).strip()
            entry = pending.setdefault(no, {"no": no, "ja": None, "zh": None})
            if is_ja(body) and entry["ja"] is None:
                entry["ja"] = body
            elif entry["zh"] is None:
                entry["zh"] = body
            else:  # 想定外の3行目は日本語/中国語の空いている方に入れる
                key = "ja" if entry["ja"] is None else "zh"
                entry[key] = body
        else:
            if is_ja(line) and intro_ja is None:
                intro_ja = line
            elif intro_zh is None:
                intro_zh = line

    roster = [pending[k] for k in sorted(pending)]
    return intro_ja, intro_zh, roster


def parse_script_cell(lines):
    """「日：」「中：」でセグメント分割。それ以外の『X：』はキュー（note）扱い"""
    segments = []
    cur = None

    def flush():
        nonlocal cur
        if cur and cur["lines"]:
            segments.append(cur)
        cur = None

    for line in lines:
        m = re.match(r"^(日|中|MC|[^：:\s]{1,14})\s*[：:]\s*(.*)$", line)
        if m:
            marker, rest = m.group(1), m.group(2).strip()
            if marker == "日":
                flush()
                cur = {"speaker": "日本語", "type": "mc1", "lines": []}
                if rest:
                    cur["lines"].append(rest)
                continue
            if marker == "中":
                flush()
                cur = {"speaker": "中文", "type": "mc2", "lines": []}
                if rest:
                    cur["lines"].append(rest)
                continue
            # MC： / 武内さん： などのキュー行
            flush()
            segments.append({"speaker": marker, "type": "cue",
                             "lines": [rest] if rest else []})
            continue

        # （…）だけの行 = ト書き（直前のセグメントには含めない）
        if re.match(r"^[（(].*[）)]$", line):
            flush()
            segments.append({"speaker": "", "type": "note", "lines": [line]})
            continue

        if cur is None:
            segments.append({"speaker": "", "type": "note", "lines": [line]})
        else:
            cur["lines"].append(line)

    flush()
    # 本文が空のキュー行も残す（「MC：」単独など）→ 表示側で見出しになる
    return [s for s in segments if s["lines"] or s["type"] == "cue"]


def extract_time(text):
    m = re.search(r"(\d{1,2}):(\d{2})", text)
    return m.group(0) if m else None


def main():
    doc = docx.Document(str(DOCX_PATH))
    table = doc.tables[0]

    parts = []
    part_no = 0

    for row in table.rows:
        cells = row.cells
        texts = [c.text.strip() for c in cells]

        # ヘッダー行
        if "時間" in texts[0] and "司会者進行" in texts[0]:
            continue

        # 結合セル = セクション見出し
        if len(set(texts)) == 1:
            title = texts[0].strip("【】 ")
            part_no += 1
            parts.append({
                "id": f"part{part_no}",
                "no": part_no,
                "title": title,
                "titleZh": SECTION_ZH.get(title, title),
                "scheduledTime": None,
                "rows": [],
            })
            continue

        if not parts:
            continue

        time_text = texts[0]
        lines = cell_lines(cells[1])
        part = parts[-1]
        if part["scheduledTime"] is None:
            part["scheduledTime"] = extract_time(time_text)

        row_obj = {"time": time_text, "script": [], "roster": None}

        # 企業紹介セクションは日中交互フォーマット
        if part["title"] == "企業様の紹介":
            intro_ja, intro_zh, roster = parse_roster(lines)
            if intro_ja:
                row_obj["script"].append({"speaker": "日本語", "type": "mc1", "lines": [intro_ja]})
            if intro_zh:
                row_obj["script"].append({"speaker": "中文", "type": "mc2", "lines": [intro_zh]})
            row_obj["roster"] = roster
        else:
            row_obj["script"] = parse_script_cell(lines)

        part["rows"].append(row_obj)

    data = dict(EVENT)
    data["parts"] = parts

    OUT_PATH.parent.mkdir(exist_ok=True)
    with open(OUT_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    # file:// でも開けるよう JS 形式でも出力（index.html はこちらを読む）
    with open(OUT_JS_PATH, "w", encoding="utf-8") as f:
        f.write("/* 自動生成: parse_script.py — 直接編集しないこと */\n")
        f.write("window.SCRIPT_DATA = ")
        json.dump(data, f, ensure_ascii=False, indent=2)
        f.write(";\n")

    total_rows = sum(len(p["rows"]) for p in parts)
    rosters = sum(len(r["roster"] or []) for p in parts for r in p["rows"])
    print(f"OK → {OUT_PATH}")
    print(f"OK → {OUT_JS_PATH}")
    print(f"   {len(parts)} sections / {total_rows} rows / roster {rosters} companies")
    for p in parts:
        segs = sum(len(r["script"]) for r in p["rows"])
        print(f"   - {p['id']} {p['title']} ({p['scheduledTime']}) segments={segs}")


if __name__ == "__main__":
    main()
