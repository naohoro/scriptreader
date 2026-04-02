#!/usr/bin/env python3
"""
parse_scripts.py
JWWA 2026 MC構成台本 (.docx) → data/*.json 変換スクリプト

使用方法: python3 parse_scripts.py
"""
import docx
import json
import re
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent  # 構成台本フォルダ
OUT_DIR  = Path(__file__).parent / "data"
OUT_DIR.mkdir(exist_ok=True)

EVENTS = [
    {
        "id": "zenjasai",
        "file": "2026 ワールドオークション前夜祭_MC構成台本.docx",
        "title": "前夜祭",
        "subtitle": "JAPANESE WAGYU WORLD AUCTION 2026",
        "date": "2026年4月3日（金）",
        "startTime": "17:00",
        "endTime": "20:30",
        "venue": "ホテル日航姫路　光琳の間（3F）",
        "mc1": {"name": "Justin", "fullName": "Justin Patterson", "lang": "英語"},
        "mc2": {"name": "Marika", "fullName": "Marika Watanabe", "lang": "日本語補足"},
        "mc1_keys": ["Justin"],
        "mc2_keys": ["Marika"],
        "both_keys": ["BOTH"],
    },
    {
        "id": "toujitsu",
        "file": "2026 ワールドオークション当日_MC構成台本.docx",
        "title": "オークション当日",
        "subtitle": "JAPANESE WAGYU WORLD AUCTION 2026",
        "date": "2026年4月4日（土）",
        "startTime": "10:00",
        "endTime": "12:00",
        "venue": "和牛マスター株式会社",
        "mc1": {"name": "Marika", "fullName": "Marika Watanabe", "lang": "英語実況 / 日本語"},
        "mc2": None,
        "mc1_keys": ["Marika"],
        "mc2_keys": [],
        "both_keys": [],
    },
    {
        "id": "hanami",
        "file": "2026 ワールドオークション花見の宴_MC構成台本.docx",
        "title": "花見の宴",
        "subtitle": "JAPANESE WAGYU WORLD AUCTION 2026",
        "date": "2026年4月4日（土）",
        "startTime": "12:30",
        "endTime": "14:30",
        "venue": "姫路市和牛マスター食肉センター",
        "mc1": {"name": "中川", "fullName": "中川", "lang": "日本語"},
        "mc2": {"name": "Horota", "fullName": "Nao Horota", "lang": "英語"},
        "mc1_keys": ["中川"],
        "mc2_keys": ["Horota"],
        "both_keys": [],
    },
]


def classify_speaker(name, event):
    """Return 'mc1', 'mc2', 'both', or 'note'"""
    if not name:
        return "note"
    for key in event.get("mc1_keys", []):
        if key in name:
            return "mc1"
    for key in event.get("mc2_keys", []):
        if key in name:
            return "mc2"
    for key in event.get("both_keys", []):
        if key in name:
            return "both"
    return "note"


def parse_mc_cell(cell, event):
    """Parse MC announcement cell → list of {speaker, type, lines}"""
    segments = []
    cur_speaker = None
    cur_type = None
    cur_lines = []

    for para in cell.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m = re.match(r"^【(.+?)】(.*)$", text)
        if m:
            # flush previous
            if cur_lines or cur_speaker is not None:
                segments.append({"speaker": cur_speaker or "", "type": cur_type or "note", "lines": cur_lines})
            cur_speaker = m.group(1).strip()
            cur_type = classify_speaker(cur_speaker, event)
            cur_lines = []
            rest = m.group(2).strip()
            if rest:
                cur_lines.append(rest)
        else:
            if cur_speaker is not None:
                cur_lines.append(text)
            else:
                # pre-speaker note
                segments.append({"speaker": "", "type": "note", "lines": [text]})

    if cur_lines or cur_speaker is not None:
        segments.append({"speaker": cur_speaker or "", "type": cur_type or "note", "lines": cur_lines})

    return [s for s in segments if s["lines"]]


def parse_action_cell(cell):
    lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "\n".join(lines)


def is_header_row(row):
    texts = [c.text.strip() for c in row.cells[:min(3, len(row.cells))]]
    return "時刻" in texts[0] and any("進行" in t or "MC" in t or "アナウンス" in t for t in texts[1:])


def is_merged_header_row(row):
    if len(row.cells) < 2:
        return False
    texts = [c.text.strip() for c in row.cells]
    non_empty = [t for t in texts if t]
    return len(non_empty) >= 2 and len(set(non_empty)) == 1


def extract_scheduled_time(title):
    m = re.search(r"(\d{1,2}:\d{2})\s*[—–-]?\s*$", title)
    if m:
        return m.group(1)
    m = re.search(r"(\d{1,2}:\d{2})", title)
    if m:
        return m.group(1)
    return None


def parse_event(cfg):
    fpath = BASE_DIR / cfg["file"]
    doc = docx.Document(str(fpath))
    parts = []

    tables = list(doc.tables)
    tbl_idx = 0

    def is_part_header_table(tbl):
        if len(tbl.rows) == 1:
            t = tbl.rows[0].cells[0].text.strip()
            return bool(re.match(r"^PART\s*\d+", t))
        return False

    def is_script_table(tbl):
        if len(tbl.columns) >= 3 and tbl.rows:
            first = [c.text.strip() for c in tbl.rows[0].cells[:3]]
            return any("時刻" in t for t in first)
        return False

    body = doc.element.body
    para_i = 0
    tbl_i = 0

    for child in body:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            text = "".join(n.text or "" for n in child.iter() if n.tag.endswith("}t")).strip()
            if re.match(r"^PART\s*\d+", text):
                num_m = re.search(r"PART\s*(\d+)", text)
                pid = f"part{num_m.group(1)}" if num_m else f"part{len(parts)+1}"
                parts.append({
                    "id": pid,
                    "title": text,
                    "scheduledTime": extract_scheduled_time(text),
                    "rows": []
                })

        elif tag == "tbl":
            tbl = tables[tbl_i]
            tbl_i += 1

            if is_part_header_table(tbl):
                title_text = tbl.rows[0].cells[0].text.strip()
                num_m = re.search(r"PART\s*(\d+)", title_text)
                pid = f"part{num_m.group(1)}" if num_m else f"part{len(parts)+1}"
                parts.append({
                    "id": pid,
                    "title": title_text,
                    "scheduledTime": extract_scheduled_time(title_text),
                    "rows": []
                })

            elif is_script_table(tbl):
                rows = []
                for row in tbl.rows:
                    if is_header_row(row):
                        continue
                    if is_merged_header_row(row):
                        continue
                    cells = row.cells
                    if len(cells) < 3:
                        continue
                    time_t = cells[0].text.strip()
                    action_t = parse_action_cell(cells[1])
                    script_segs = parse_mc_cell(cells[2], cfg)
                    if time_t or action_t or script_segs:
                        rows.append({"time": time_t, "action": action_t, "script": script_segs})

                if rows:
                    if parts:
                        parts[-1]["rows"].extend(rows)
                    else:
                        parts.append({"id": "part0", "title": "Opening", "scheduledTime": None, "rows": rows})

    return {
        "id": cfg["id"],
        "title": cfg["title"],
        "subtitle": cfg["subtitle"],
        "date": cfg["date"],
        "startTime": cfg["startTime"],
        "endTime": cfg["endTime"],
        "venue": cfg["venue"],
        "mc1": cfg["mc1"],
        "mc2": cfg["mc2"],
        "parts": parts,
    }


if __name__ == "__main__":
    for cfg in EVENTS:
        print(f"Processing: {cfg['file']} ...", end=" ")
        data = parse_event(cfg)
        out = OUT_DIR / f"{cfg['id']}.json"
        with open(out, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        total_rows = sum(len(p["rows"]) for p in data["parts"])
        print(f"OK → {out.name}  ({len(data['parts'])} parts, {total_rows} rows)")
    print("Done.")
